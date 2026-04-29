"""
Export generators - ZERO Django imports.
build_excel() and build_docx() return bytes.
"""
import io
import re

from processing.review_annotations import build_review_indexes, collect_legacy_cell_flags


# helpers
_NAME_HEADERS = re.compile(r"\b(name|employee|payee|employer|member)\b", re.IGNORECASE)
_MONEY_HEADER_RE = re.compile(
    r"\b(amount|pay|salary|total|gross|net|balance|value|price|cost|tax|fee|deduction|contribution|nssf|nhif|paye)\b",
    re.IGNORECASE,
)
_CURRENCY_MARKER_RE = re.compile(r"[$\u20ac\u00a3\u00a5]|(?:USD|UGX|KES|TZS|RWF|EUR|GBP)\b", re.IGNORECASE)
_LEADING_CURRENCY_CODE_RE = re.compile(r"^(?:USD|UGX|KES|TZS|RWF|EUR|GBP)\s*", re.IGNORECASE)
_TRAILING_CURRENCY_CODE_RE = re.compile(r"\s*(?:USD|UGX|KES|TZS|RWF|EUR|GBP)$", re.IGNORECASE)
_EDGE_SYMBOL_RE = re.compile(r"^[\s=#$%&/\u00a3\u20ac\u00a5]+|[\s=#$%&/\u00a3\u20ac\u00a5]+$")
_NUMERIC_RE = re.compile(r"^-?[\d,]+(?:\.\d+)?$")


def _clean_money(val: str) -> str:
    """Strip leading/trailing symbols/currency markers from numeric money values."""
    raw = (val or "").strip()
    if not raw:
        return ""
    cleaned = _LEADING_CURRENCY_CODE_RE.sub("", raw)
    cleaned = _TRAILING_CURRENCY_CODE_RE.sub("", cleaned)
    cleaned = _EDGE_SYMBOL_RE.sub("", cleaned).strip()
    return cleaned if _NUMERIC_RE.match(cleaned) else raw


def _is_money_col(header: str) -> bool:
    """Heuristic: column header mentions amount/pay/salary/total."""
    return bool(_MONEY_HEADER_RE.search((header or "").strip()))


def _looks_money_col(values: list[str]) -> bool:
    for val in values:
        raw = (val or "").strip()
        if not raw:
            continue
        if _CURRENCY_MARKER_RE.search(raw):
            return True
        cleaned = _clean_money(raw)
        if cleaned != raw and _NUMERIC_RE.match(cleaned):
            return True
    return False


def _is_name_col(header: str) -> bool:
    return bool(_NAME_HEADERS.search(header or ""))


def _cell_dict(cell, row_idx=None, col_idx=None) -> dict:
    if isinstance(cell, dict):
        cell_data = dict(cell)
    else:
        cell_data = {"content": str(cell or "")}
    cell_data.setdefault("raw_content", str(cell_data.get("content", "") or ""))
    if row_idx is not None:
        cell_data.setdefault("row_index", row_idx)
    if col_idx is not None:
        cell_data.setdefault("column_index", col_idx)
    return cell_data


def _cell_text(cell) -> str:
    return str(_cell_dict(cell).get("content", "") or "")


def _process_table(rows: list) -> list:
    """Clean money columns and capitalize name columns while preserving cell metadata."""
    if not rows:
        return rows

    headers = [str(_cell_dict(cell).get("raw_content", "")) for cell in (rows[0] if rows else [])]
    money_cols = {i for i, h in enumerate(headers) if _is_money_col(h)}
    for i in range(len(headers)):
        if i in money_cols:
            continue
        col_values = [
            str(_cell_dict(row[i]).get("raw_content", ""))
            for row in rows
            if len(row) > i
        ]
        if _looks_money_col(col_values):
            money_cols.add(i)
    name_cols = {i for i, h in enumerate(headers) if _is_name_col(h)}

    result = []
    for row_idx, row in enumerate(rows):
        new_row = []
        for i, cell in enumerate(row):
            cell_data = _cell_dict(cell, row_idx=row_idx, col_idx=i)
            val = str(cell_data.get("raw_content", cell_data.get("content", "")))
            if i in money_cols:
                val = _clean_money(val)
            if i in name_cols and row_idx > 0 and val:
                val = val.upper()
            cell_data["content"] = val
            new_row.append(cell_data)
        result.append(new_row)
    return result


def _add_highlighted_text(para, text: str, flagged_words: list, font_name: str = "Courier New", font_pt: int = 9):
    """Add runs to *para* with flagged word occurrences highlighted in yellow."""
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt

    if not text:
        return
    unique = sorted({w for w in flagged_words if w}, key=len, reverse=True)
    if not unique:
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_pt)
        return
    pattern = "(" + "|".join(re.escape(w) for w in unique) + ")"
    flagged_set = set(flagged_words)
    for part in re.split(pattern, text):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(font_pt)
        if part in flagged_set:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_highlighted_text_by_ranges(para, text: str, ranges: list, font_name: str = "Courier New", font_pt: int = 9):
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt

    if not text:
        return

    merged = []
    seen = set()
    for item in ranges or []:
        offset = item.get("span_offset")
        length = item.get("span_length")
        if offset is None or not length:
            continue
        try:
            start = int(offset)
            end = start + int(length)
        except (TypeError, ValueError):
            continue
        if end <= start:
            continue
        key = (start, end)
        if key not in seen:
            merged.append(key)
            seen.add(key)

    merged.sort()
    cursor = 0
    for start, end in merged:
        if start > cursor:
            run = para.add_run(text[cursor:start])
            run.font.name = font_name
            run.font.size = Pt(font_pt)
        if end > cursor:
            run = para.add_run(text[max(cursor, start):end])
            run.font.name = font_name
            run.font.size = Pt(font_pt)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            cursor = end

    if cursor < len(text):
        run = para.add_run(text[cursor:])
        run.font.name = font_name
        run.font.size = Pt(font_pt)


def _safe_sheet_title(title: str, existing_titles: list[str]) -> str:
    cleaned = re.sub(r"[:\\/?*\[\]]", "_", title)[:31] or "Sheet"
    if cleaned not in existing_titles:
        return cleaned
    suffix = 2
    while True:
        candidate = f"{cleaned[:28]}_{suffix}"
        if candidate not in existing_titles:
            return candidate
        suffix += 1


def _review_fill(flags, low_fill, med_fill):
    min_conf = min(float(flag.get("confidence", 1.0)) for flag in flags)
    return low_fill if min_conf < 0.5 else med_fill


def _review_comment(page_number, table_index, row_number, column_number, flags):
    lines = [
        "Needs manual review",
        f"Page {page_number}, Table {table_index}, Row {row_number}, Column {column_number}",
    ]
    for flag in flags:
        lines.append(f'- {flag.get("word", "").strip() or "token"} ({float(flag.get("confidence", 0.0)) * 100:.1f}%)')
    return "\n".join(lines)


# Excel

def build_excel(pages: list, review_flags: list) -> bytes:
    """
    Build an Excel workbook:
    - One sheet per extracted table ("P1_Table_1", "P2_Table_1", ...)
    - A "Flagged_for_Review" sheet for review flags
    Returns bytes.
    """
    from openpyxl import Workbook
    from openpyxl.comments import Comment
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    _, cell_flags_by_page = build_review_indexes(review_flags, unresolved_only=True)

    wb = Workbook()
    wb.remove(wb.active)  # remove default sheet

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="0F1A2E")
    alt_fill = PatternFill("solid", fgColor="F8FAFC")
    low_fill = PatternFill("solid", fgColor="FEE2E2")   # red-100
    med_fill = PatternFill("solid", fgColor="FEF3C7")   # amber-100

    for page_data in pages or []:
        page_number = page_data.get("page_number")
        for table_idx, table in enumerate(page_data.get("tables") or [], start=1):
            title = _safe_sheet_title(f"P{page_number}_Table_{table_idx}", wb.sheetnames)
            ws = wb.create_sheet(title=title)
            ws.freeze_panes = "A2"
            processed = _process_table(table)
            for row_idx, row in enumerate(processed, start=1):
                for col_idx, cell_data in enumerate(row, start=1):
                    val = _cell_text(cell_data)
                    cell = ws.cell(row=row_idx, column=col_idx, value=val)
                    if row_idx == 1:
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal="center")
                    elif row_idx % 2 == 0:
                        cell.fill = alt_fill

                    flags = list(
                        cell_flags_by_page.get(page_number, {}).get((table_idx, row_idx - 1, col_idx - 1), [])
                    )
                    if not flags:
                        flags = collect_legacy_cell_flags(review_flags, page_number, val, unresolved_only=True)
                    if flags:
                        cell.fill = _review_fill(flags, low_fill, med_fill)
                        cell.comment = Comment(
                            _review_comment(page_number, table_idx, row_idx, col_idx, flags),
                            "Oryx",
                        )
                        if row_idx == 1:
                            cell.font = Font(bold=True, color="111827")
                        else:
                            cell.font = Font(color="111827")

            for col in ws.columns:
                max_len = max((len(str(c.value or "")) for c in col), default=0)
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)

    # Flagged for review sheet
    ws_flags = wb.create_sheet(title="Flagged_for_Review")
    flag_headers = ["Word", "Confidence", "Page", "Table", "Row", "Column", "Reviewed", "Corrected Value"]
    for ci, h in enumerate(flag_headers, start=1):
        cell = ws_flags.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = PatternFill("solid", fgColor="B45309")  # amber

    for ri, flag in enumerate(review_flags, start=2):
        page_num = flag.page_result.page_number if hasattr(flag, "page_result") and flag.page_result else ""
        conf_fill = low_fill if float(flag.confidence) < 0.5 else med_fill
        word_cell = ws_flags.cell(row=ri, column=1, value=flag.word)
        word_cell.fill = conf_fill
        conf_cell = ws_flags.cell(row=ri, column=2, value=round(flag.confidence * 100, 1))
        conf_cell.fill = conf_fill
        ws_flags.cell(row=ri, column=3, value=page_num)
        ws_flags.cell(row=ri, column=4, value=getattr(flag, "table_index", "") or "")
        ws_flags.cell(row=ri, column=5, value=(getattr(flag, "row_index", None) + 1) if getattr(flag, "row_index", None) is not None else "")
        ws_flags.cell(row=ri, column=6, value=(getattr(flag, "column_index", None) + 1) if getattr(flag, "column_index", None) is not None else "")
        ws_flags.cell(row=ri, column=7, value="Yes" if flag.reviewed else "No")
        ws_flags.cell(row=ri, column=8, value=flag.corrected_value or "")

    for col in ws_flags.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=0)
        ws_flags.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# Word

def build_docx(pages: list, extracted_fields: dict, review_flags: list) -> bytes:
    """
    Build a Word document:
    - Title page (Oryx)
    - Per-page: extracted text + tables
    - Section: flagged words highlighted in red
    Returns bytes.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    text_ranges_by_page, cell_flags_by_page = build_review_indexes(review_flags, unresolved_only=True)

    doc = Document()

    # Title page
    title_para = doc.add_heading("Oryx - Extraction Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if extracted_fields:
        doc.add_heading("Extracted Fields", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"
        for k, v in extracted_fields.items():
            row = table.add_row().cells
            row[0].text = k.replace("_", " ").title()
            row[1].text = str(v)
        doc.add_paragraph()

    # Per-page content
    for page_data in pages:
        doc.add_heading(f"Page {page_data['page_number']}", level=2)

        text = page_data.get("extracted_text", "")
        if text:
            para = doc.add_paragraph()
            page_ranges = text_ranges_by_page.get(page_data["page_number"], [])
            if page_ranges:
                _add_highlighted_text_by_ranges(para, text, page_ranges)
            else:
                _add_highlighted_text(para, text, [])

        # Tables
        for table_idx, tbl_data in enumerate(page_data.get("tables", []), start=1):
            if not tbl_data:
                continue
            processed = _process_table(tbl_data)
            num_cols = max(len(row) for row in processed) if processed else 0
            if num_cols == 0:
                continue
            t = doc.add_table(rows=len(processed), cols=num_cols)
            t.style = "Table Grid"
            for ri, row in enumerate(processed):
                for ci, cell_data in enumerate(row):
                    cell = t.cell(ri, ci)
                    value = _cell_text(cell_data)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    cell_flags = list(
                        cell_flags_by_page.get(page_data["page_number"], {}).get((table_idx, ri, ci), [])
                    )
                    if not cell_flags:
                        cell_flags = collect_legacy_cell_flags(
                            review_flags,
                            page_data["page_number"],
                            value,
                            unresolved_only=True,
                        )
                    if cell_flags:
                        _add_highlighted_text(para, value, [flag["word"] for flag in cell_flags], font_name="Calibri", font_pt=10)
                    else:
                        para.add_run(value)
                    if ri == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
            doc.add_paragraph()

        doc.add_page_break()

    # Flagged words
    if review_flags:
        doc.add_heading("Flagged for Review", level=1)
        for flag in review_flags:
            para = doc.add_paragraph()
            label = para.add_run(
                f"[Page {flag.page_result.page_number if hasattr(flag, 'page_result') and flag.page_result else '?'}] "
            )
            label.font.color.rgb = RGBColor(0x71, 0x71, 0x71)
            word_run = para.add_run(flag.word)
            word_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # red-600
            word_run.bold = True
            conf_run = para.add_run(f"  ({flag.confidence:.0%} confidence)")
            conf_run.font.size = Pt(9)
            conf_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
